"""
Document parser for extracting text from various file formats.
Optimized for handling large documents (100+ pages).
"""
import logging
import os
from typing import Dict, Any, List, Optional, Generator
import io
import math

# PDF extraction
from pypdf import PdfReader

# DOCX extraction
import docx

# For error handling
from slideforge.core.exceptions import ProcessingError

# Setup logging
logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parser for extracting text and metadata from different document formats.
    Supports PDF, DOCX, and TXT files with optimizations for large documents.
    """
    
    # Default chunk size in pages for large documents
    DEFAULT_CHUNK_SIZE = 10
    
    # Default maximum characters to extract (approximately 100 pages of text)
    MAX_CHARS = 250000
    
    @staticmethod
    def parse(file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Parse a document file and extract text and metadata.
        
        Args:
            file_path: Path to the document file
            file_type: Type of the document file (pdf, docx, txt)
            
        Returns:
            Dict containing extracted text and metadata
            
        Raises:
            ProcessingError: If there is an error during extraction
        """
        file_type = file_type.lower()
        
        try:
            if file_type == "pdf":
                return DocumentParser._parse_pdf(file_path)
            elif file_type == "docx":
                return DocumentParser._parse_docx(file_path)
            elif file_type == "txt":
                return DocumentParser._parse_txt(file_path)
            else:
                raise ProcessingError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to parse {file_type} file: {str(e)}")

    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """
        Parse a PDF file and extract text and metadata.
        Optimized for large documents with intelligent chunking.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dict containing extracted text and metadata
        """
        logger.info(f"Parsing PDF file: {file_path}")
        
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            logger.info(f"PDF has {total_pages} pages")
            
            # Extract metadata
            metadata = reader.metadata
            if metadata:
                metadata_dict = {
                    "title": metadata.get("/Title", ""),
                    "author": metadata.get("/Author", ""),
                    "subject": metadata.get("/Subject", ""),
                    "keywords": metadata.get("/Keywords", ""),
                    "creator": metadata.get("/Creator", ""),
                    "producer": metadata.get("/Producer", ""),
                    "creation_date": str(metadata.get("/CreationDate", "")),
                }
            else:
                metadata_dict = {}
            
            # For large documents, use intelligent extraction
            if total_pages > 30:  # Consider anything over 30 pages as a large document
                return DocumentParser._extract_large_pdf(reader, metadata_dict, file_path)
            
            # For smaller documents, extract all text
            text = ""
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"Page {page_num + 1}:\n{page_text}\n\n"
            
            # If text extraction failed, report it
            if not text.strip():
                logger.warning(f"Failed to extract text from PDF: {file_path}")
                text = "[No extractable text found in the PDF document]"
            
            return {
                "text": text,
                "metadata": metadata_dict,
                "pages": total_pages,
                "file_path": file_path,
                "file_type": "pdf",
                "is_large_document": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to parse PDF file: {str(e)}")
    
    @staticmethod
    def _extract_large_pdf(reader: PdfReader, metadata_dict: Dict[str, Any], file_path: str) -> Dict[str, Any]:
        """
        Extract text from a large PDF using intelligent chunking.
        
        Args:
            reader: PdfReader object
            metadata_dict: Extracted metadata
            file_path: Path to the PDF file
            
        Returns:
            Dict containing extracted text and metadata
        """
        total_pages = len(reader.pages)
        logger.info(f"Using large document extraction for {total_pages}-page PDF")
        
        # Extract TOC, intro, and conclusion if possible
        important_sections = {
            "introduction": "",
            "conclusion": "",
            "table_of_contents": ""
        }
        
        # Extract table of contents (usually in the first few pages)
        for i in range(min(5, total_pages)):
            page_text = reader.pages[i].extract_text()
            if page_text and ("content" in page_text.lower() or "table of" in page_text.lower()):
                important_sections["table_of_contents"] = page_text
                logger.info(f"Table of contents extracted from page {i+1}")
                break
        
        # Extract introduction (usually in the first 10% of the document)
        intro_limit = min(max(3, int(total_pages * 0.1)), 10)  # 3 to 10 pages
        intro_text = ""
        for i in range(intro_limit):
            if i >= total_pages:
                break
            page_text = reader.pages[i].extract_text()
            if page_text:
                intro_text += page_text + "\n\n"
        important_sections["introduction"] = intro_text
        
        # Extract conclusion (usually in the last 10% of the document)
        conclusion_limit = min(max(3, int(total_pages * 0.1)), 10)  # 3 to 10 pages
        conclusion_text = ""
        for i in range(max(0, total_pages - conclusion_limit), total_pages):
            if i >= total_pages:
                break
            page_text = reader.pages[i].extract_text()
            if page_text:
                conclusion_text += page_text + "\n\n"
        important_sections["conclusion"] = conclusion_text
        
        # Extract strategically distributed chunks
        chunks = []
        
        # Define chunk distribution logic (beginning, middle, end, and strategic points)
        if total_pages <= 50:
            # For medium documents (30-50 pages), take chunks from beginning, middle, and end
            chunk_points = [0, total_pages // 2, max(0, total_pages - 10)]
            chunk_size = 5  # 5 pages per chunk
        elif total_pages <= 100:
            # For larger documents (50-100 pages), take more distributed chunks
            chunk_points = [
                0,  # Beginning
                total_pages // 4,  # First quarter
                total_pages // 2,  # Middle
                (total_pages * 3) // 4,  # Third quarter
                max(0, total_pages - 10)  # End
            ]
            chunk_size = 3  # 3 pages per chunk
        else:
            # For very large documents (100+ pages), take more sparse chunks
            num_chunks = min(10, total_pages // 20)  # Up to 10 chunks, at least 20 pages apart
            chunk_points = [int((total_pages * i) / num_chunks) for i in range(num_chunks)]
            chunk_points.append(max(0, total_pages - 5))  # Always include the end
            chunk_size = 2  # 2 pages per chunk
        
        # Extract text from each chunk point
        for start_page in chunk_points:
            chunk_text = ""
            for i in range(start_page, min(start_page + chunk_size, total_pages)):
                page_text = reader.pages[i].extract_text()
                if page_text:
                    chunk_text += f"Page {i+1}:\n{page_text}\n\n"
            if chunk_text:
                chunks.append(chunk_text)
        
        # Combine important sections and chunks
        combined_text = ""
        
        # Add metadata-based context
        title = metadata_dict.get("title", "")
        if title:
            combined_text += f"DOCUMENT TITLE: {title}\n\n"
        
        # Add table of contents if available
        if important_sections["table_of_contents"]:
            combined_text += "TABLE OF CONTENTS:\n" + important_sections["table_of_contents"] + "\n\n"
        
        # Add introduction
        if important_sections["introduction"]:
            combined_text += "INTRODUCTION:\n" + important_sections["introduction"] + "\n\n"
        
        # Add content samples from throughout the document
        combined_text += "CONTENT SAMPLES FROM THROUGHOUT THE DOCUMENT:\n\n"
        for i, chunk in enumerate(chunks):
            combined_text += f"--- CONTENT SAMPLE {i+1} ---\n{chunk}\n\n"
        
        # Add conclusion
        if important_sections["conclusion"]:
            combined_text += "CONCLUSION:\n" + important_sections["conclusion"] + "\n\n"
        
        # Ensure we don't exceed a reasonable size for LLM processing
        if len(combined_text) > DocumentParser.MAX_CHARS:
            logger.warning(f"Large document extracted text exceeds {DocumentParser.MAX_CHARS} characters, truncating")
            combined_text = combined_text[:DocumentParser.MAX_CHARS] + "...[CONTENT TRUNCATED DUE TO SIZE]..."
        
        return {
            "text": combined_text,
            "metadata": metadata_dict,
            "pages": total_pages,
            "file_path": file_path,
            "file_type": "pdf",
            "is_large_document": True,
            "extracted_sections": {
                "has_toc": bool(important_sections["table_of_contents"]),
                "has_intro": bool(important_sections["introduction"]),
                "has_conclusion": bool(important_sections["conclusion"]),
                "num_chunks": len(chunks)
            }
        }

    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text and metadata.
        Optimized for large documents.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dict containing extracted text and metadata
        """
        logger.info(f"Parsing DOCX file: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            prop = doc.core_properties
            metadata_dict = {
                "title": prop.title or "",
                "author": prop.author or "",
                "subject": prop.subject or "",
                "keywords": prop.keywords or "",
                "category": prop.category or "",
                "comments": prop.comments or "",
                "created": str(prop.created) if prop.created else "",
                "modified": str(prop.modified) if prop.modified else "",
            }
            
            # Count paragraphs to determine document size
            total_paragraphs = len(doc.paragraphs)
            logger.info(f"DOCX has {total_paragraphs} paragraphs")
            
            # For large documents, use intelligent extraction
            if total_paragraphs > 500:  # Consider anything over 500 paragraphs as a large document
                return DocumentParser._extract_large_docx(doc, metadata_dict, file_path)
            
            # Extract paragraphs
            text = ""
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\nTable:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
                text += "\n"
            
            # If text extraction failed, report it
            if not text.strip():
                logger.warning(f"Failed to extract text from DOCX: {file_path}")
                text = "[No extractable text found in the DOCX document]"
            
            return {
                "text": text,
                "metadata": metadata_dict,
                "paragraphs": total_paragraphs,
                "file_path": file_path,
                "file_type": "docx",
                "is_large_document": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to parse DOCX file: {str(e)}")
    
    @staticmethod
    def _extract_large_docx(doc: docx.Document, metadata_dict: Dict[str, Any], file_path: str) -> Dict[str, Any]:
        """
        Extract text from a large DOCX using intelligent chunking.
        
        Args:
            doc: docx.Document object
            metadata_dict: Extracted metadata
            file_path: Path to the DOCX file
            
        Returns:
            Dict containing extracted text and metadata
        """
        total_paragraphs = len(doc.paragraphs)
        logger.info(f"Using large document extraction for {total_paragraphs}-paragraph DOCX")
        
        # Extract document structure by analyzing headings
        structure = []
        
        # Extract the table of contents (look for paragraphs with "Contents" or "Table of Contents")
        toc_paragraphs = []
        found_toc = False
        for i, para in enumerate(doc.paragraphs[:50]):  # Look only in the first 50 paragraphs
            if not found_toc and para.text and ("content" in para.text.lower() or "table of" in para.text.lower()):
                found_toc = True
            
            if found_toc:
                toc_paragraphs.append(para.text)
                # Stop when we find what appears to be the end of the TOC
                if len(toc_paragraphs) > 5 and not para.text.strip():
                    break
        
        # Extract heading structure
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
                headings.append({
                    "level": level,
                    "text": para.text
                })
        
        # Extract introduction (first 10% of paragraphs)
        intro_limit = min(max(20, int(total_paragraphs * 0.1)), 100)  # 20 to 100 paragraphs
        intro_text = ""
        for i in range(intro_limit):
            if i >= len(doc.paragraphs):
                break
            intro_text += doc.paragraphs[i].text + "\n"
        
        # Extract conclusion (last 10% of paragraphs)
        conclusion_limit = min(max(20, int(total_paragraphs * 0.1)), 100)  # 20 to 100 paragraphs
        conclusion_text = ""
        for i in range(max(0, total_paragraphs - conclusion_limit), total_paragraphs):
            if i >= len(doc.paragraphs):
                break
            conclusion_text += doc.paragraphs[i].text + "\n"
        
        # Extract strategically distributed chunks
        chunks = []
        
        # Define chunk distribution logic based on document size
        if total_paragraphs <= 1000:
            # For medium-large documents (500-1000 paragraphs)
            chunk_points = [0, total_paragraphs // 4, total_paragraphs // 2, 
                          (total_paragraphs * 3) // 4, max(0, total_paragraphs - 50)]
            chunk_size = 30  # 30 paragraphs per chunk
        else:
            # For very large documents (1000+ paragraphs)
            num_chunks = min(10, total_paragraphs // 200)  # Up to 10 chunks
            chunk_points = [int((total_paragraphs * i) / num_chunks) for i in range(num_chunks)]
            chunk_points.append(max(0, total_paragraphs - 50))  # Always include the end
            chunk_size = 20  # 20 paragraphs per chunk
        
        # Extract text from each chunk point
        for start_para in chunk_points:
            chunk_text = ""
            for i in range(start_para, min(start_para + chunk_size, total_paragraphs)):
                if i < len(doc.paragraphs):
                    chunk_text += doc.paragraphs[i].text + "\n"
            if chunk_text:
                chunks.append(chunk_text)
        
        # Combine into a comprehensive representation of the document
        combined_text = ""
        
        # Add metadata-based context
        title = metadata_dict.get("title", "")
        if title:
            combined_text += f"DOCUMENT TITLE: {title}\n\n"
        
        # Add document structure overview
        if headings:
            combined_text += "DOCUMENT STRUCTURE:\n"
            for heading in headings[:30]:  # Include up to 30 headings
                indent = "  " * (heading["level"] - 1)
                combined_text += f"{indent}- {heading['text']}\n"
            if len(headings) > 30:
                combined_text += "  [Additional headings omitted for brevity]\n"
            combined_text += "\n"
        
        # Add table of contents if available
        if toc_paragraphs:
            combined_text += "TABLE OF CONTENTS:\n" + "\n".join(toc_paragraphs) + "\n\n"
        
        # Add introduction
        combined_text += "INTRODUCTION:\n" + intro_text + "\n\n"
        
        # Add content samples from throughout the document
        combined_text += "CONTENT SAMPLES FROM THROUGHOUT THE DOCUMENT:\n\n"
        for i, chunk in enumerate(chunks):
            combined_text += f"--- CONTENT SAMPLE {i+1} ---\n{chunk}\n\n"
        
        # Add conclusion
        combined_text += "CONCLUSION:\n" + conclusion_text + "\n\n"
        
        # Ensure we don't exceed a reasonable size for LLM processing
        if len(combined_text) > DocumentParser.MAX_CHARS:
            logger.warning(f"Large document extracted text exceeds {DocumentParser.MAX_CHARS} characters, truncating")
            combined_text = combined_text[:DocumentParser.MAX_CHARS] + "...[CONTENT TRUNCATED DUE TO SIZE]..."
        
        return {
            "text": combined_text,
            "metadata": metadata_dict,
            "paragraphs": total_paragraphs,
            "file_path": file_path,
            "file_type": "docx",
            "is_large_document": True,
            "extracted_sections": {
                "has_toc": bool(toc_paragraphs),
                "has_headings": bool(headings),
                "has_intro": bool(intro_text),
                "has_conclusion": bool(conclusion_text),
                "num_chunks": len(chunks)
            }
        }

    @staticmethod
    def _parse_txt(file_path: str) -> Dict[str, Any]:
        """
        Parse a TXT file and extract text.
        Optimized for large files.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Dict containing extracted text
        """
        logger.info(f"Parsing TXT file: {file_path}")
        
        try:
            # Get file size
            file_size = os.path.getsize(file_path)
            logger.info(f"TXT file size: {file_size / 1024:.2f} KB")
            
            # For large files, use chunked reading
            if file_size > 1024 * 1024:  # 1 MB
                return DocumentParser._extract_large_txt(file_path)
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # If all encodings fail, use binary and decode with replacement
                with open(file_path, 'rb') as f:
                    content = f.read()
                    text = content.decode('utf-8', errors='replace')
            
            # Split into lines for basic structure
            lines = text.splitlines()
            
            # Extract basic metadata (if available in a structured format)
            metadata_dict = {}
            
            # Try to get title from first non-empty line
            for line in lines:
                if line.strip():
                    metadata_dict["title"] = line.strip()
                    break
            
            return {
                "text": text,
                "metadata": metadata_dict,
                "lines": len(lines),
                "file_path": file_path,
                "file_type": "txt",
                "is_large_document": False
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to parse TXT file: {str(e)}")
    
    @staticmethod
    def _extract_large_txt(file_path: str) -> Dict[str, Any]:
        """
        Extract text from a large TXT file using chunking.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Dict containing extracted text
        """
        logger.info(f"Using large document extraction for TXT file: {file_path}")
        
        try:
            # Get file size
            file_size = os.path.getsize(file_path)
            
            # Determine encoding by trying different encodings on the first few KB
            encoding = 'utf-8'  # Default
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            with open(file_path, 'rb') as f:
                sample = f.read(min(file_size, 4096))  # Read at most 4 KB
                
                for enc in encodings_to_try:
                    try:
                        sample.decode(enc)
                        encoding = enc
                        break
                    except UnicodeDecodeError:
                        continue
            
            # Extract beginning, some middle chunks, and end
            beginning = ""
            middle_chunks = []
            end = ""
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                # Read the first 1000 lines
                lines = []
                for _ in range(1000):
                    line = f.readline()
                    if not line:
                        break
                    lines.append(line)
                
                beginning = "".join(lines)
                
                # Count total lines by reading through file
                total_lines = len(lines)
                while f.readline():
                    total_lines += 1
                
                logger.info(f"TXT file has approximately {total_lines} lines")
                
                # Reset file position
                f.seek(0)
                
                # Extract strategically distributed chunks
                if total_lines > 5000:
                    # For very large text files
                    chunk_points = [
                        total_lines // 4,       # 25%
                        total_lines // 2,       # 50%
                        (total_lines * 3) // 4  # 75%
                    ]
                    chunk_size = 200  # 200 lines per chunk
                else:
                    # For moderately large text files
                    chunk_points = [
                        total_lines // 3,       # 33%
                        (total_lines * 2) // 3  # 66%
                    ]
                    chunk_size = 300  # 300 lines per chunk
                
                # Extract middle chunks
                for start_line in chunk_points:
                    # Skip to the start line
                    f.seek(0)
                    for _ in range(min(start_line, total_lines)):
                        f.readline()
                    
                    # Read chunk_size lines
                    chunk_lines = []
                    for _ in range(chunk_size):
                        line = f.readline()
                        if not line:
                            break
                        chunk_lines.append(line)
                    
                    if chunk_lines:
                        middle_chunks.append("".join(chunk_lines))
                
                # Extract the end (last 1000 lines)
                f.seek(0)
                if total_lines > 1000:
                    for _ in range(total_lines - 1000):
                        f.readline()
                
                end_lines = f.readlines()
                end = "".join(end_lines)
            
            # Combine chunks
            combined_text = ""
            combined_text += "BEGINNING OF DOCUMENT:\n" + beginning + "\n\n"
            
            for i, chunk in enumerate(middle_chunks):
                combined_text += f"MIDDLE SECTION {i+1}:\n{chunk}\n\n"
            
            combined_text += "END OF DOCUMENT:\n" + end
            
            # Extract basic metadata
            metadata_dict = {}
            
            # Try to get title from first non-empty line
            for line in beginning.splitlines():
                if line.strip():
                    metadata_dict["title"] = line.strip()
                    break
            
            # Ensure we don't exceed a reasonable size for LLM processing
            if len(combined_text) > DocumentParser.MAX_CHARS:
                logger.warning(f"Large document extracted text exceeds {DocumentParser.MAX_CHARS} characters, truncating")
                combined_text = combined_text[:DocumentParser.MAX_CHARS] + "...[CONTENT TRUNCATED DUE TO SIZE]..."
            
            return {
                "text": combined_text,
                "metadata": metadata_dict,
                "lines": total_lines,
                "file_path": file_path,
                "file_type": "txt",
                "is_large_document": True,
                "extracted_sections": {
                    "has_beginning": bool(beginning),
                    "num_middle_chunks": len(middle_chunks),
                    "has_end": bool(end)
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting large TXT file {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to parse large TXT file: {str(e)}")